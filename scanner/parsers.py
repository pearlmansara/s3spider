"""
parsers.py - Extract plain text content from various file types.
Supports: plain text, .docx, .xlsx, .pdf, and all config/code file types.
"""

import io
import logging

logger = logging.getLogger("s3spider")

# ── Extension groups ──────────────────────────────────────────────────────────

TEXT_EXTENSIONS = {
    ".txt", ".log", ".csv", ".env", ".conf", ".config", ".yaml", ".yml",
    ".json", ".xml", ".ini", ".cfg", ".sh", ".bash", ".zsh", ".py", ".rb",
    ".js", ".ts", ".php", ".sql", ".md", ".toml", ".properties", ".gradle",
    ".tf", ".tfvars", ".tfstate", ".tfstate.backup",
    ".pem", ".key", ".crt", ".cer", ".pub",
}

DOCX_EXTENSIONS = {".docx", ".doc"}
XLSX_EXTENSIONS = {".xlsx", ".xls"}
PDF_EXTENSIONS  = {".pdf"}

ALL_SUPPORTED_EXTENSIONS = (
    TEXT_EXTENSIONS | DOCX_EXTENSIONS | XLSX_EXTENSIONS | PDF_EXTENSIONS
)


def is_supported(key: str) -> bool:
    """Return True if the S3 object key has a supported extension."""
    lower = key.lower()
    return any(lower.endswith(ext) for ext in ALL_SUPPORTED_EXTENSIONS)


def extract_text(data: bytes, key: str) -> str | None:
    """
    Given raw bytes and the S3 object key, return extracted plain text.
    Returns None if the file type is unsupported or extraction fails.
    """
    lower = key.lower()

    if any(lower.endswith(ext) for ext in PDF_EXTENSIONS):
        return _extract_pdf(data)

    if any(lower.endswith(ext) for ext in DOCX_EXTENSIONS):
        return _extract_docx(data)

    if any(lower.endswith(ext) for ext in XLSX_EXTENSIONS):
        return _extract_xlsx(data)

    if any(lower.endswith(ext) for ext in TEXT_EXTENSIONS):
        return _extract_text(data)

    return None


# ── Internal extractors ───────────────────────────────────────────────────────

def _extract_text(data: bytes) -> str | None:
    """Decode bytes as UTF-8 text (with fallback to latin-1)."""
    try:
        return data.decode("utf-8", errors="replace")
    except Exception as e:
        logger.debug(f"Text decode error: {e}")
        try:
            return data.decode("latin-1", errors="replace")
        except Exception:
            return None


def _extract_pdf(data: bytes) -> str | None:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    lines.append(text)
        return "\n".join(lines) if lines else ""
    except ImportError:
        logger.warning("pdfplumber not installed — skipping PDF files. Install with: pip install pdfplumber")
        return None
    except Exception as e:
        logger.debug(f"PDF extraction error: {e}")
        return None


def _extract_docx(data: bytes) -> str | None:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Also grab table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — skipping .docx files. Install with: pip install python-docx")
        return None
    except Exception as e:
        logger.debug(f"DOCX extraction error: {e}")
        return None


def _extract_xlsx(data: bytes) -> str | None:
    """Extract text from a .xlsx file using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    except ImportError:
        logger.warning("openpyxl not installed — skipping .xlsx files. Install with: pip install openpyxl")
        return None
    except Exception as e:
        logger.debug(f"XLSX extraction error: {e}")
        return None
